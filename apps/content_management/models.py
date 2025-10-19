"""
Modelos para el módulo de Gestión de Contenido Publicitario
Sistema PubliTrack - Gestión de cuñas publicitarias y archivos de audio
INCLUYE: Sistema de Contratos con Plantillas Automáticas
"""

import os
import uuid
from decimal import Decimal
from datetime import datetime, timedelta
from django.db import models
from django.conf import settings
from django.core.validators import MinValueValidator, MaxValueValidator, FileExtensionValidator
from django.utils import timezone
from django.urls import reverse
from django.core.exceptions import ValidationError
from mutagen import File as MutagenFile
from mutagen.mp3 import MP3
from mutagen.wave import WAVE
from mutagen.mp4 import MP4

# ==================== FUNCIONES AUXILIARES ====================

def audio_upload_path(instance, filename):
    """
    Genera la ruta de subida para archivos de audio
    """
    now = timezone.now()
    year = now.year
    month = now.strftime('%m')
    ext = filename.split('.')[-1].lower()
    unique_filename = f"{uuid.uuid4().hex}.{ext}"
    return f"audio_spots/{year}/{month}/{unique_filename}"


def contract_template_path(instance, filename):
    """Genera la ruta de subida para plantillas de contrato"""
    ext = filename.split('.')[-1].lower()
    unique_filename = f"template_{uuid.uuid4().hex}.{ext}"
    return f"contract_templates/{unique_filename}"


def contract_output_path(instance, filename):
    """Genera la ruta de subida para contratos generados"""
    year = timezone.now().year
    month = timezone.now().strftime('%m')
    ext = filename.split('.')[-1].lower()
    unique_filename = f"contrato_{instance.numero_contrato}_{uuid.uuid4().hex[:8]}.{ext}"
    return f"contracts/{year}/{month}/{unique_filename}"


def numero_a_letras(numero):
    """
    Convierte un número decimal a su representación en letras (español)
    Ejemplo: 180.50 -> "CIENTO OCHENTA CON 50/100 DÓLARES AMERICANOS"
    """
    UNIDADES = ['', 'UNO', 'DOS', 'TRES', 'CUATRO', 'CINCO', 'SEIS', 'SIETE', 'OCHO', 'NUEVE']
    DECENAS = ['', 'DIEZ', 'VEINTE', 'TREINTA', 'CUARENTA', 'CINCUENTA', 
               'SESENTA', 'SETENTA', 'OCHENTA', 'NOVENTA']
    ESPECIALES = ['DIEZ', 'ONCE', 'DOCE', 'TRECE', 'CATORCE', 'QUINCE',
                  'DIECISÉIS', 'DIECISIETE', 'DIECIOCHO', 'DIECINUEVE']
    CENTENAS = ['', 'CIENTO', 'DOSCIENTOS', 'TRESCIENTOS', 'CUATROCIENTOS',
                'QUINIENTOS', 'SEISCIENTOS', 'SETECIENTOS', 'OCHOCIENTOS', 'NOVECIENTOS']
    
    def convertir_grupo(n):
        """Convierte un grupo de hasta 3 dígitos"""
        if n == 0:
            return ''
        elif n == 100:
            return 'CIEN'
        
        centena = n // 100
        decena = (n % 100) // 10
        unidad = n % 10
        
        resultado = []
        
        if centena > 0:
            resultado.append(CENTENAS[centena])
        
        resto = n % 100
        if 10 <= resto <= 19:
            resultado.append(ESPECIALES[resto - 10])
        else:
            if decena > 0:
                if decena == 2 and unidad > 0:
                    resultado.append('VEINTI' + UNIDADES[unidad])
                    return ' '.join(resultado)
                else:
                    resultado.append(DECENAS[decena])
            if unidad > 0:
                if decena > 0 and decena != 2:
                    resultado.append('Y')
                resultado.append(UNIDADES[unidad])
        
        return ' '.join(resultado)
    
    # Separar parte entera y decimal
    if isinstance(numero, (int, float)):
        numero = Decimal(str(numero))
    
    partes = str(numero).split('.')
    parte_entera = int(partes[0])
    parte_decimal = partes[1][:2].ljust(2, '0') if len(partes) > 1 else '00'
    
    if parte_entera == 0:
        resultado = 'CERO'
    elif parte_entera < 1000:
        resultado = convertir_grupo(parte_entera)
    elif parte_entera < 1000000:
        miles = parte_entera // 1000
        resto = parte_entera % 1000
        
        if miles == 1:
            resultado = 'MIL'
        else:
            resultado = convertir_grupo(miles) + ' MIL'
        
        if resto > 0:
            resultado += ' ' + convertir_grupo(resto)
    else:
        millones = parte_entera // 1000000
        resto = parte_entera % 1000000
        
        if millones == 1:
            resultado = 'UN MILLÓN'
        else:
            resultado = convertir_grupo(millones) + ' MILLONES'
        
        if resto > 0:
            if resto >= 1000:
                miles = resto // 1000
                resto_final = resto % 1000
                if miles > 0:
                    resultado += ' ' + convertir_grupo(miles) + ' MIL'
                if resto_final > 0:
                    resultado += ' ' + convertir_grupo(resto_final)
            else:
                resultado += ' ' + convertir_grupo(resto)
    
    # Formato final
    return f"{resultado} CON {parte_decimal}/100 DÓLARES AMERICANOS"


# ==================== MODELOS EXISTENTES ====================

class CategoriaPublicitaria(models.Model):
    """
    Categorías para clasificar el tipo de publicidad
    """
    
    nombre = models.CharField(
        'Nombre de la Categoría',
        max_length=100,
        unique=True,
        help_text='Nombre de la categoría publicitaria'
    )
    
    descripcion = models.TextField(
        'Descripción',
        blank=True,
        help_text='Descripción detallada de la categoría'
    )
    
    color_codigo = models.CharField(
        'Código de Color',
        max_length=7,
        default='#007bff',
        help_text='Color para identificar la categoría (formato hex)'
    )
    
    tarifa_base = models.DecimalField(
        'Tarifa Base',
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        help_text='Tarifa base por segundo para esta categoría'
    )
    
    is_active = models.BooleanField(
        'Activa',
        default=True,
        help_text='Si la categoría está activa para nuevas cuñas'
    )
    
    created_at = models.DateTimeField('Creada', auto_now_add=True)
    updated_at = models.DateTimeField('Actualizada', auto_now=True)
    
    class Meta:
        verbose_name = 'Categoría Publicitaria'
        verbose_name_plural = 'Categorías Publicitarias'
        ordering = ['nombre']
    
    def __str__(self):
        return self.nombre
    
    def get_absolute_url(self):
        return reverse('content:categoria_detail', kwargs={'pk': self.pk})


class TipoContrato(models.Model):
    """
    Tipos de contrato publicitario (paquetes, por tiempo, etc.)
    """
    
    DURACION_CHOICES = [
        ('semanal', 'Semanal'),
        ('mensual', 'Mensual'),
        ('trimestral', 'Trimestral'),
        ('semestral', 'Semestral'),
        ('anual', 'Anual'),
        ('personalizado', 'Personalizado'),
    ]
    
    nombre = models.CharField(
        'Nombre del Tipo',
        max_length=100,
        unique=True
    )
    
    descripcion = models.TextField(
        'Descripción',
        blank=True
    )
    
    duracion_tipo = models.CharField(
        'Tipo de Duración',
        max_length=20,
        choices=DURACION_CHOICES,
        default='mensual'
    )
    
    duracion_dias = models.PositiveIntegerField(
        'Duración en Días',
        default=30,
        help_text='Duración estándar del contrato en días'
    )
    
    repeticiones_minimas = models.PositiveIntegerField(
        'Repeticiones Mínimas',
        default=1,
        help_text='Número mínimo de repeticiones por día'
    )
    
    descuento_porcentaje = models.DecimalField(
        'Descuento (%)',
        max_digits=5,
        decimal_places=2,
        default=Decimal('0.00'),
        validators=[MinValueValidator(0), MaxValueValidator(100)],
        help_text='Descuento aplicable a este tipo de contrato'
    )
    
    is_active = models.BooleanField('Activo', default=True)
    created_at = models.DateTimeField('Creado', auto_now_add=True)
    updated_at = models.DateTimeField('Actualizado', auto_now=True)
    
    class Meta:
        verbose_name = 'Tipo de Contrato'
        verbose_name_plural = 'Tipos de Contrato'
        ordering = ['nombre']
    
    def __str__(self):
        return f"{self.nombre} ({self.get_duracion_tipo_display()})"


class ArchivoAudio(models.Model):
    """
    Archivo de audio de la cuña publicitaria
    """
    
    FORMATO_CHOICES = [
        ('mp3', 'MP3'),
        ('wav', 'WAV'),
        ('aac', 'AAC'),
        ('m4a', 'M4A'),
        ('ogg', 'OGG'),
    ]
    
    CALIDAD_CHOICES = [
        ('baja', 'Baja (128 kbps)'),
        ('media', 'Media (192 kbps)'),
        ('alta', 'Alta (320 kbps)'),
        ('sin_perdida', 'Sin Pérdida'),
    ]
    
    archivo = models.FileField(
        'Archivo de Audio',
        upload_to=audio_upload_path,
        validators=[FileExtensionValidator(allowed_extensions=['mp3', 'wav', 'aac', 'm4a', 'ogg'])],
        help_text='Archivo de audio de la cuña publicitaria'
    )
    
    nombre_original = models.CharField(
        'Nombre Original',
        max_length=255,
        help_text='Nombre original del archivo subido'
    )
    
    formato = models.CharField(
        'Formato',
        max_length=10,
        choices=FORMATO_CHOICES,
        help_text='Formato del archivo de audio'
    )
    
    duracion_segundos = models.PositiveIntegerField(
        'Duración (segundos)',
        null=True,
        blank=True,
        help_text='Duración real del audio en segundos'
    )
    
    tamaño_bytes = models.PositiveBigIntegerField(
        'Tamaño (bytes)',
        null=True,
        blank=True,
        help_text='Tamaño del archivo en bytes'
    )
    
    bitrate = models.PositiveIntegerField(
        'Bitrate (kbps)',
        null=True,
        blank=True,
        help_text='Bitrate del audio en kbps'
    )
    
    sample_rate = models.PositiveIntegerField(
        'Sample Rate (Hz)',
        null=True,
        blank=True,
        help_text='Frecuencia de muestreo en Hz'
    )
    
    canales = models.PositiveIntegerField(
        'Canales',
        null=True,
        blank=True,
        help_text='Número de canales de audio (1=mono, 2=estéreo)'
    )
    
    calidad = models.CharField(
        'Calidad',
        max_length=20,
        choices=CALIDAD_CHOICES,
        null=True,
        blank=True,
        help_text='Calidad estimada del audio'
    )
    
    hash_archivo = models.CharField(
        'Hash del Archivo',
        max_length=64,
        unique=True,
        null=True,
        blank=True,
        help_text='Hash SHA256 del archivo para detectar duplicados'
    )
    
    metadatos_extra = models.JSONField(
        'Metadatos Adicionales',
        default=dict,
        blank=True,
        help_text='Metadatos adicionales extraídos del archivo'
    )
    
    subido_por = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        related_name='archivos_audio_subidos',
        verbose_name='Subido por'
    )
    
    fecha_subida = models.DateTimeField('Fecha de Subida', auto_now_add=True)
    
    class Meta:
        verbose_name = 'Archivo de Audio'
        verbose_name_plural = 'Archivos de Audio'
        ordering = ['-fecha_subida']
    
    def __str__(self):
        return f"{self.nombre_original} ({self.duracion_formateada})"
    
    def save(self, *args, **kwargs):
        """Override save para extraer metadatos automáticamente"""
        if self.archivo and not self.duracion_segundos:
            self.extraer_metadatos()
        super().save(*args, **kwargs)
    
    def extraer_metadatos(self):
        """
        Extrae metadatos del archivo de audio usando mutagen
        """
        try:
            if self.archivo and os.path.exists(self.archivo.path):
                self.tamaño_bytes = os.path.getsize(self.archivo.path)
                self.nombre_original = os.path.basename(self.archivo.name)
                
                ext = os.path.splitext(self.archivo.name)[1].lower().lstrip('.')
                self.formato = ext if ext in dict(self.FORMATO_CHOICES) else 'mp3'
                
                audio_file = MutagenFile(self.archivo.path)
                
                if audio_file is not None:
                    if hasattr(audio_file, 'info') and hasattr(audio_file.info, 'length'):
                        self.duracion_segundos = int(audio_file.info.length)
                    
                    if hasattr(audio_file, 'info') and hasattr(audio_file.info, 'bitrate'):
                        self.bitrate = audio_file.info.bitrate
                        
                        if self.bitrate:
                            if self.bitrate < 160:
                                self.calidad = 'baja'
                            elif self.bitrate < 256:
                                self.calidad = 'media'
                            elif self.bitrate < 400:
                                self.calidad = 'alta'
                            else:
                                self.calidad = 'sin_perdida'
                    
                    if hasattr(audio_file, 'info') and hasattr(audio_file.info, 'sample_rate'):
                        self.sample_rate = audio_file.info.sample_rate
                    
                    if hasattr(audio_file, 'info') and hasattr(audio_file.info, 'channels'):
                        self.canales = audio_file.info.channels
                    
                    if audio_file.tags:
                        self.metadatos_extra = {
                            'title': str(audio_file.tags.get('TIT2', [''])[0]) if audio_file.tags.get('TIT2') else '',
                            'artist': str(audio_file.tags.get('TPE1', [''])[0]) if audio_file.tags.get('TPE1') else '',
                            'album': str(audio_file.tags.get('TALB', [''])[0]) if audio_file.tags.get('TALB') else '',
                        }
                
        except Exception as e:
            print(f"Error extrayendo metadatos: {e}")
    
    @property
    def duracion_formateada(self):
        """Retorna la duración en formato mm:ss"""
        if self.duracion_segundos:
            minutos = self.duracion_segundos // 60
            segundos = self.duracion_segundos % 60
            return f"{minutos:02d}:{segundos:02d}"
        return "00:00"
    
    @property
    def tamaño_formateado(self):
        """Retorna el tamaño en formato legible"""
        if self.tamaño_bytes:
            if self.tamaño_bytes < 1024:
                return f"{self.tamaño_bytes} B"
            elif self.tamaño_bytes < 1024**2:
                return f"{self.tamaño_bytes/1024:.1f} KB"
            elif self.tamaño_bytes < 1024**3:
                return f"{self.tamaño_bytes/(1024**2):.1f} MB"
            else:
                return f"{self.tamaño_bytes/(1024**3):.1f} GB"
        return "0 B"
    
    def get_absolute_url(self):
        return reverse('content:audio_detail', kwargs={'pk': self.pk})


class CuñaPublicitaria(models.Model):
    """
    Modelo principal para las cuñas publicitarias
    """
    
    ESTADO_CHOICES = [
        ('borrador', 'Borrador'),
        ('pendiente_revision', 'Pendiente de Revisión'),
        ('en_produccion', 'En Producción'),
        ('aprobada', 'Aprobada'),
        ('activa', 'Activa'),
        ('pausada', 'Pausada'),
        ('finalizada', 'Finalizada'),
        ('cancelada', 'Cancelada'),
    ]
    
    PRIORIDAD_CHOICES = [
        ('baja', 'Baja'),
        ('normal', 'Normal'),
        ('alta', 'Alta'),
        ('urgente', 'Urgente'),
    ]
    
    # Información básica
    codigo = models.CharField(
        'Código',
        max_length=20,
        unique=True,
        help_text='Código único de la cuña (generado automáticamente)'
    )
    
    titulo = models.CharField(
        'Título',
        max_length=200,
        help_text='Título descriptivo de la cuña publicitaria'
    )
    
    descripcion = models.TextField(
        'Descripción',
        blank=True,
        help_text='Descripción detallada del contenido publicitario'
    )
    
    # Relaciones
    cliente = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.CASCADE,
        limit_choices_to={'rol': 'cliente'},
        related_name='cuñas_publicitarias',
        verbose_name='Cliente'
    )
    
    vendedor_asignado = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        limit_choices_to={'rol': 'vendedor'},
        related_name='cuñas_vendidas',
        verbose_name='Vendedor Asignado'
    )
    
    categoria = models.ForeignKey(
        CategoriaPublicitaria,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='cuñas',
        verbose_name='Categoría'
    )
    
    tipo_contrato = models.ForeignKey(
        TipoContrato,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='cuñas',
        verbose_name='Tipo de Contrato'
    )
    
    archivo_audio = models.ForeignKey(
        ArchivoAudio,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='cuñas',
        verbose_name='Archivo de Audio'
    )
    
    # Estado y prioridad
    estado = models.CharField(
        'Estado',
        max_length=20,
        choices=ESTADO_CHOICES,
        default='borrador',
        help_text='Estado actual de la cuña'
    )
    
    prioridad = models.CharField(
        'Prioridad',
        max_length=10,
        choices=PRIORIDAD_CHOICES,
        default='normal',
        help_text='Prioridad de la cuña'
    )
    
    # Información técnica
    duracion_planeada = models.PositiveIntegerField(
        'Duración Planeada (segundos)',
        help_text='Duración planeada de la cuña en segundos'
    )
    
    # Información comercial
    precio_total = models.DecimalField(
        'Precio Total',
        max_digits=10,
        decimal_places=2,
        help_text='Precio total del contrato publicitario'
    )
    
    precio_por_segundo = models.DecimalField(
        'Precio por Segundo',
        max_digits=8,
        decimal_places=4,
        null=True,
        blank=True,
        help_text='Precio por segundo de duración'
    )
    
    excluir_sabados = models.BooleanField(
        'Excluir Sábados',
        default=False,
        help_text='Excluir sábados de la programación'
    )

    excluir_domingos = models.BooleanField(
        'Excluir Domingos', 
        default=False,
        help_text='Excluir domingos de la programación'
    )

    repeticiones_dia = models.PositiveIntegerField(
        'Repeticiones por Día',
        default=1,
        help_text='Número de veces que se reproduce por día'
    )
    
    # Fechas del contrato
    fecha_inicio = models.DateField(
        'Fecha de Inicio',
        help_text='Fecha de inicio de la campaña publicitaria'
    )
    
    fecha_fin = models.DateField(
        'Fecha de Fin',
        help_text='Fecha de finalización de la campaña publicitaria'
    )
    
    # Información de aprobación
    aprobada_por = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='cuñas_aprobadas',
        verbose_name='Aprobada por'
    )
    
    fecha_aprobacion = models.DateTimeField(
        'Fecha de Aprobación',
        null=True,
        blank=True
    )
    
    observaciones = models.TextField(
        'Observaciones',
        blank=True,
        help_text='Observaciones sobre la cuña publicitaria'
    )
    
    # Configuraciones adicionales
    requiere_aprobacion = models.BooleanField(
        'Requiere Aprobación',
        default=True,
        help_text='Si la cuña requiere aprobación antes de ser transmitida'
    )
    
    permite_edicion = models.BooleanField(
        'Permite Edición',
        default=True,
        help_text='Si la cuña puede ser editada después de creada'
    )
    
    notificar_vencimiento = models.BooleanField(
        'Notificar Vencimiento',
        default=True,
        help_text='Si notificar cuando esté próxima a vencer'
    )
    
    dias_aviso_vencimiento = models.PositiveIntegerField(
        'Días de Aviso de Vencimiento',
        default=7,
        help_text='Días antes del vencimiento para enviar notificación'
    )
    
    # Tags y palabras clave
    tags = models.CharField(
        'Tags',
        max_length=500,
        blank=True,
        help_text='Palabras clave separadas por comas para búsqueda'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        related_name='cuñas_creadas',
        verbose_name='Creada por'
    )
    
    created_at = models.DateTimeField('Creada', auto_now_add=True)
    updated_at = models.DateTimeField('Actualizada', auto_now=True)
    
    class Meta:
        verbose_name = 'Cuña Publicitaria'
        verbose_name_plural = 'Cuñas Publicitarias'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['estado', 'fecha_inicio']),
            models.Index(fields=['cliente', 'estado']),
            models.Index(fields=['vendedor_asignado', 'estado']),
            models.Index(fields=['fecha_inicio', 'fecha_fin']),
        ]
    
    def __str__(self):
        return f"{self.codigo} - {self.titulo}"
    
    def save(self, *args, **kwargs):
        """Override save para generar código automáticamente"""
        if not self.codigo:
            self.codigo = self.generar_codigo()
        
        if not self.precio_por_segundo and self.duracion_planeada:
            self.precio_por_segundo = self.precio_total / self.duracion_planeada
        
        super().save(*args, **kwargs)
    
    def generar_codigo(self):
        """Genera un código único para la cuña"""
        año = timezone.now().year
        mes = timezone.now().month
        
        count = CuñaPublicitaria.objects.filter(
            created_at__year=año,
            created_at__month=mes
        ).count() + 1
        
        return f"CP{año}{mes:02d}{count:04d}"
    
    def clean(self):
        """Validaciones personalizadas"""
        if self.fecha_inicio and self.fecha_fin:
            if self.fecha_fin <= self.fecha_inicio:
                raise ValidationError('La fecha de fin debe ser posterior a la fecha de inicio.')
        
        if self.archivo_audio and self.archivo_audio.duracion_segundos:
            diferencia = abs(self.duracion_planeada - self.archivo_audio.duracion_segundos)
            if diferencia > 5:
                raise ValidationError(
                    f'La duración planeada ({self.duracion_planeada}s) difiere mucho '
                    f'del archivo de audio ({self.archivo_audio.duracion_segundos}s).'
                )
    
    @property
    def dias_efectivos(self):
        """Calcula los días efectivos considerando exclusiones"""
        if not self.fecha_inicio or not self.fecha_fin:
            return 0
    
        from datetime import timedelta
        dias_totales = 0
        fecha_actual = self.fecha_inicio
    
        while fecha_actual <= self.fecha_fin:
            if self.excluir_sabados and fecha_actual.weekday() == 5:
                pass
            elif self.excluir_domingos and fecha_actual.weekday() == 6:
                pass
            else:
                dias_totales += 1
        
            fecha_actual += timedelta(days=1)
    
        return dias_totales

    @property
    def emisiones_totales_reales(self):
        """Calcula el total real de emisiones considerando exclusiones"""
        return self.dias_efectivos * self.repeticiones_dia

    @property
    def precio_total_calculado(self):
        """Calcula el precio total correctamente"""
        if not self.precio_por_segundo:
            return Decimal('0.00')
        return Decimal(str(
            self.duracion_planeada * 
            self.repeticiones_dia * 
            float(self.precio_por_segundo) * 
            self.dias_efectivos
        ))

    @property
    def costo_por_emision_real(self):
        """Calcula el costo por emisión individual considerando días efectivos"""
        emisiones_totales = self.emisiones_totales_reales
        if emisiones_totales > 0:
            return self.precio_total / emisiones_totales
        return Decimal('0.00')
    
    @property
    def dias_restantes(self):
        """Calcula los días restantes hasta el fin de la campaña"""
        if self.fecha_fin:
            hoy = timezone.now().date()
            if self.fecha_fin >= hoy:
                return (self.fecha_fin - hoy).days
            return 0
        return None
    
    @property
    def esta_activa(self):
        """Verifica si la cuña está en periodo activo"""
        hoy = timezone.now().date()
        return (
            self.estado == 'activa' and
            self.fecha_inicio <= hoy <= self.fecha_fin
        )
    
    @property
    def esta_vencida(self):
        """Verifica si la cuña está vencida"""
        hoy = timezone.now().date()
        return self.fecha_fin < hoy if self.fecha_fin else False
    
    @property
    def requiere_notificacion_vencimiento(self):
        """Verifica si debe notificar por próximo vencimiento"""
        if not self.notificar_vencimiento or not self.fecha_fin:
            return False
        
        hoy = timezone.now().date()
        dias_hasta_vencimiento = (self.fecha_fin - hoy).days
        
        return (
            0 <= dias_hasta_vencimiento <= self.dias_aviso_vencimiento and
            self.estado in ['activa', 'aprobada']
        )
    
    @property
    def duracion_total_dias(self):
        """Calcula la duración total de la campaña en días"""
        if self.fecha_inicio and self.fecha_fin:
            return (self.fecha_fin - self.fecha_inicio).days + 1
        return 0
    
    @property
    def reproducciones_totales(self):
        """Calcula el total de reproducciones planeadas"""
        return self.duracion_total_dias * self.repeticiones_dia
    
    @property
    def costo_por_reproduccion(self):
        """Calcula el costo por cada reproducción"""
        if self.reproducciones_totales > 0:
            return self.precio_total / self.reproducciones_totales
        return Decimal('0.00')
    
    @property
    def semaforo_estado(self):
        """Determina color de semáforo basado en estado y fechas"""
        if self.esta_vencida:
            return 'rojo'
        elif self.requiere_notificacion_vencimiento:
            return 'amarillo'
        elif self.estado in ['activa', 'aprobada']:
            return 'verde'
        elif self.estado in ['cancelada', 'finalizada']:
            return 'gris'
        else:
            return 'amarillo'
    
    def aprobar(self, user):
        """Aprueba la cuña publicitaria"""
        self.estado = 'aprobada'
        self.aprobada_por = user
        self.fecha_aprobacion = timezone.now()
        self.save()
    
    def activar(self):
        """Activa la cuña publicitaria"""
        if self.estado == 'aprobada':
            self.estado = 'activa'
            self.save()
    
    def pausar(self):
        """Pausa la cuña publicitaria"""
        if self.estado == 'activa':
            self.estado = 'pausada'
            self.save()
    
    def finalizar(self):
        """Finaliza la cuña publicitaria"""
        self.estado = 'finalizada'
        self.save()
    
    def get_absolute_url(self):
        return reverse('content:cuña_detail', kwargs={'pk': self.pk})


class HistorialCuña(models.Model):
    """
    Historial de cambios de las cuñas publicitarias
    """
    
    ACCION_CHOICES = [
        ('creada', 'Creada'),
        ('editada', 'Editada'),
        ('aprobada', 'Aprobada'),
        ('activada', 'Activada'),
        ('pausada', 'Pausada'),
        ('finalizada', 'Finalizada'),
        ('cancelada', 'Cancelada'),
        ('audio_subido', 'Audio Subido'),
        ('audio_cambiado', 'Audio Cambiado'),
    ]
    
    cuña = models.ForeignKey(
        CuñaPublicitaria,
        on_delete=models.CASCADE,
        related_name='historial',
        verbose_name='Cuña'
    )
    
    accion = models.CharField(
        'Acción',
        max_length=20,
        choices=ACCION_CHOICES
    )
    
    usuario = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Usuario'
    )
    
    descripcion = models.TextField(
        'Descripción',
        help_text='Descripción detallada del cambio realizado'
    )
    
    datos_anteriores = models.JSONField(
        'Datos Anteriores',
        null=True,
        blank=True,
        help_text='Estado anterior de los campos modificados'
    )
    
    datos_nuevos = models.JSONField(
        'Datos Nuevos',
        null=True,
        blank=True,
        help_text='Nuevo estado de los campos modificados'
    )
    
    fecha = models.DateTimeField('Fecha', auto_now_add=True)
    
    class Meta:
        verbose_name = 'Historial de Cuña'
        verbose_name_plural = 'Historial de Cuñas'
        ordering = ['-fecha']
    
    def __str__(self):
        return f"{self.cuña.codigo} - {self.get_accion_display()} - {self.fecha.strftime('%d/%m/%Y %H:%M')}"


# ==================== MODELOS DE CONTRATOS ====================

class PlantillaContrato(models.Model):
    """
    Plantillas de contrato reutilizables en formato Word
    Permite subir diferentes modelos de contrato que se rellenarán automáticamente
    """
    
    TIPO_CONTRATO_CHOICES = [
        ('publicidad_tv', 'Publicidad Televisión'),
        ('publicidad_radio', 'Publicidad Radio'),
        ('paquete_mensual', 'Paquete Mensual'),
        ('paquete_trimestral', 'Paquete Trimestral'),
        ('paquete_semestral', 'Paquete Semestral'),
        ('contrato_especial', 'Contrato Especial'),
        ('otro', 'Otro'),
    ]
    
    nombre = models.CharField(
        'Nombre de la Plantilla',
        max_length=200,
        help_text='Nombre descriptivo (ej: Contrato Publicidad TV 2025)'
    )
    
    tipo_contrato = models.CharField(
        'Tipo de Contrato',
        max_length=50,
        choices=TIPO_CONTRATO_CHOICES,
        default='publicidad_tv',
        help_text='Tipo de contrato que representa esta plantilla'
    )
    
    descripcion = models.TextField(
        'Descripción',
        blank=True,
        help_text='Descripción de cuándo usar esta plantilla'
    )
    
    archivo_plantilla = models.FileField(
        'Archivo de Plantilla (.docx)',
        upload_to=contract_template_path,
        validators=[FileExtensionValidator(allowed_extensions=['docx'])],
        help_text='Archivo Word con marcadores: {{NOMBRE_CLIENTE}}, {{RUC_DNI}}, etc.'
    )
    
    # Configuración de la plantilla
    incluye_iva = models.BooleanField(
        'Incluye IVA',
        default=True,
        help_text='Si el contrato incluye cálculo de IVA'
    )
    
    porcentaje_iva = models.DecimalField(
        'Porcentaje IVA',
        max_digits=5,
        decimal_places=2,
        default=Decimal('15.00'),
        help_text='Porcentaje de IVA a aplicar'
    )
    
    variables_disponibles = models.JSONField(
        'Variables Disponibles',
        default=dict,
        blank=True,
        help_text='Variables que se reemplazan en la plantilla'
    )
    
    # Instrucciones de uso
    instrucciones = models.TextField(
        'Instrucciones de Uso',
        blank=True,
        help_text='Instrucciones sobre cómo usar esta plantilla'
    )
    
    # Control de versiones
    version = models.CharField(
        'Versión',
        max_length=20,
        default='1.0',
        help_text='Versión de la plantilla'
    )
    
    is_active = models.BooleanField(
        'Activa',
        default=True,
        help_text='Si la plantilla está activa'
    )
    
    is_default = models.BooleanField(
        'Plantilla por Defecto',
        default=False,
        help_text='Si es la plantilla predeterminada'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        related_name='plantillas_contrato_creadas',
        verbose_name='Creada por'
    )
    
    created_at = models.DateTimeField('Creada', auto_now_add=True)
    updated_at = models.DateTimeField('Actualizada', auto_now=True)
    
    class Meta:
        verbose_name = 'Plantilla de Contrato'
        verbose_name_plural = 'Plantillas de Contrato'
        ordering = ['-is_default', '-created_at']
        indexes = [
            models.Index(fields=['tipo_contrato', 'is_active']),
            models.Index(fields=['is_default']),
        ]
    
    def __str__(self):
        return f"{self.nombre} (v{self.version})"
    
    def save(self, *args, **kwargs):
        """Override save para generar variables disponibles"""
        if not self.variables_disponibles:
            self.variables_disponibles = {
                'NOMBRE_CLIENTE': 'Nombre o Razón Social del Cliente',
                'RUC_DNI': 'RUC o Cédula del Cliente',
                'CIUDAD': 'Ciudad del Cliente',
                'PROVINCIA': 'Provincia del Cliente',
                'DIRECCION_EXACTA': 'Dirección completa del Cliente',
                'DIRECCION': 'Dirección (alias de DIRECCION_EXACTA)',
                'VALOR_NUMEROS': 'Valor del contrato en números (ej: 180.00)',
                'VALOR_LETRAS': 'Valor del contrato en letras',
                'IVA_NUMEROS': 'Valor del IVA en números',
                'IVA_LETRAS': 'Valor del IVA en letras',
                'TOTAL_NUMEROS': 'Total con IVA en números',
                'TOTAL_LETRAS': 'Total con IVA en letras',
                'FECHA_INICIO': 'Fecha de inicio del contrato',
                'FECHA_FIN': 'Fecha de fin del contrato',
                'DURACION_DIAS': 'Duración en días',
                'DURACION_MESES': 'Duración en meses',
                'SPOTS_DIA': 'Spots por día',
                'DURACION_SPOT': 'Duración del spot en segundos',
                'FECHA_ACTUAL': 'Fecha actual de generación',
                'NUMERO_CONTRATO': 'Número único del contrato',
            }
        
        # Si se marca como default, desmarcar las demás
        if self.is_default:
            PlantillaContrato.objects.filter(
                tipo_contrato=self.tipo_contrato,
                is_default=True
            ).exclude(pk=self.pk).update(is_default=False)
        
        super().save(*args, **kwargs)
    
    def get_absolute_url(self):
        return reverse('content:plantilla_contrato_detail', kwargs={'pk': self.pk})


class ContratoGenerado(models.Model):
    """
    Contratos generados a partir de plantillas
    Guarda el archivo rellenado y los datos utilizados
    """
    
    ESTADO_CHOICES = [
        ('borrador', 'Borrador'),
        ('generado', 'Generado'),
        ('enviado', 'Enviado al Cliente'),
        ('firmado', 'Firmado'),
        ('activo', 'Activo'),
        ('vencido', 'Vencido'),
        ('cancelado', 'Cancelado'),
    ]
    
    # Información básica
    numero_contrato = models.CharField(
        'Número de Contrato',
        max_length=50,
        unique=True,
        help_text='Número único del contrato (generado automáticamente)'
    )
    
    cuña = models.ForeignKey(
        CuñaPublicitaria,
        on_delete=models.CASCADE,
        related_name='contratos',
        verbose_name='Cuña Publicitaria',
        help_text='Cuña asociada a este contrato'
    )
    
    plantilla_usada = models.ForeignKey(
        PlantillaContrato,
        on_delete=models.SET_NULL,
        null=True,
        related_name='contratos_generados',
        verbose_name='Plantilla Utilizada'
    )
    
    # Cliente (denormalizado para histórico)
    cliente = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.CASCADE,
        related_name='contratos_generados',
        verbose_name='Cliente'
    )
    
    nombre_cliente = models.CharField(
        'Nombre del Cliente',
        max_length=255,
        help_text='Nombre almacenado al momento de generación'
    )
    
    ruc_dni_cliente = models.CharField(
        'RUC/DNI del Cliente',
        max_length=20,
        help_text='RUC/DNI almacenado al momento de generación'
    )
    
    # Archivo del contrato generado
    archivo_contrato = models.FileField(
        'Archivo del Contrato',
        upload_to=contract_output_path,
        blank=True,
        null=True,
        help_text='Archivo Word generado con datos del contrato'
    )
    
    archivo_contrato_pdf = models.FileField(
        'Contrato en PDF',
        upload_to=contract_output_path,
        blank=True,
        null=True,
        help_text='Versión PDF del contrato (opcional)'
    )
    
    # Valores del contrato (denormalizados)
    valor_sin_iva = models.DecimalField(
        'Valor sin IVA',
        max_digits=10,
        decimal_places=2,
        help_text='Valor del contrato sin IVA'
    )
    
    valor_iva = models.DecimalField(
        'Valor IVA',
        max_digits=10,
        decimal_places=2,
        default=Decimal('0.00'),
        help_text='Valor del IVA'
    )
    
    valor_total = models.DecimalField(
        'Valor Total',
        max_digits=10,
        decimal_places=2,
        help_text='Valor total con IVA incluido'
    )
    
    # Datos utilizados en la generación
    datos_generacion = models.JSONField(
        'Datos de Generación',
        default=dict,
        help_text='Todos los datos usados para rellenar la plantilla'
    )
    
    # Estado y fechas
    estado = models.CharField(
        'Estado',
        max_length=20,
        choices=ESTADO_CHOICES,
        default='borrador',
        help_text='Estado actual del contrato'
    )
    
    fecha_generacion = models.DateTimeField(
        'Fecha de Generación',
        auto_now_add=True
    )
    
    fecha_envio = models.DateTimeField(
        'Fecha de Envío',
        null=True,
        blank=True,
        help_text='Fecha en que se envió al cliente'
    )
    
    fecha_firma = models.DateTimeField(
        'Fecha de Firma',
        null=True,
        blank=True,
        help_text='Fecha en que se firmó el contrato'
    )
    
    # Observaciones
    observaciones = models.TextField(
        'Observaciones',
        blank=True,
        help_text='Observaciones sobre el contrato'
    )
    
    # Metadatos
    generado_por = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        related_name='contratos_generados_por',
        verbose_name='Generado por'
    )
    
    created_at = models.DateTimeField('Creado', auto_now_add=True)
    updated_at = models.DateTimeField('Actualizado', auto_now=True)
    
    class Meta:
        verbose_name = 'Contrato Generado'
        verbose_name_plural = 'Contratos Generados'
        ordering = ['-fecha_generacion']
        indexes = [
            models.Index(fields=['numero_contrato']),
            models.Index(fields=['cliente', 'estado']),
            models.Index(fields=['cuña']),
            models.Index(fields=['fecha_generacion']),
        ]
    
    def __str__(self):
        return f"Contrato {self.numero_contrato} - {self.nombre_cliente}"
    
    def save(self, *args, **kwargs):
        """Override save para generar número de contrato"""
        if not self.numero_contrato:
            self.numero_contrato = self.generar_numero_contrato()
        
        # Calcular totales si no están definidos
        if self.valor_sin_iva and not self.valor_total:
            if self.plantilla_usada and self.plantilla_usada.incluye_iva:
                porcentaje_iva = self.plantilla_usada.porcentaje_iva / 100
                self.valor_iva = self.valor_sin_iva * Decimal(str(porcentaje_iva))
                self.valor_total = self.valor_sin_iva + self.valor_iva
            else:
                self.valor_iva = Decimal('0.00')
                self.valor_total = self.valor_sin_iva
        
        super().save(*args, **kwargs)
    
    def generar_numero_contrato(self):
        """Genera un número único para el contrato"""
        año = timezone.now().year
        mes = timezone.now().month
        
        count = ContratoGenerado.objects.filter(
            fecha_generacion__year=año,
            fecha_generacion__month=mes
        ).count() + 1
        
        return f"CTR{año}{mes:02d}{count:04d}"
    
    def generar_contrato(self):
        """
        Genera el archivo de contrato rellenando la plantilla con datos del cliente y cuña
        Requiere: pip install python-docx
        """
        try:
            from docx import Document
            from io import BytesIO
            import locale
            
            # Configurar locale para fechas en español
            try:
                locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_TIME, 'es_ES')
                except:
                    pass
            
            if not self.plantilla_usada or not self.plantilla_usada.archivo_plantilla:
                raise ValueError("No hay plantilla asignada")
            
            # Abrir la plantilla
            doc = Document(self.plantilla_usada.archivo_plantilla.path)
            
            # Preparar datos para reemplazo
            cliente = self.cuña.cliente
            cuña = self.cuña
            
            # Calcular valores
            valor_sin_iva = self.valor_sin_iva or cuña.precio_total
            if self.plantilla_usada.incluye_iva:
                porcentaje_iva = self.plantilla_usada.porcentaje_iva / 100
                valor_iva = valor_sin_iva * Decimal(str(porcentaje_iva))
                valor_total = valor_sin_iva + valor_iva
            else:
                valor_iva = Decimal('0.00')
                valor_total = valor_sin_iva
            
            # Calcular duración en meses (aproximado)
            duracion_dias = cuña.duracion_total_dias
            duracion_meses = round(duracion_dias / 30, 1)
            
            # Diccionario de reemplazo
            datos_reemplazo = {
                '{{NOMBRE_CLIENTE}}': cliente.empresa or cliente.razon_social or cliente.get_full_name(),
                '{{RUC_DNI}}': cliente.ruc_dni or 'N/A',
                '{{CIUDAD}}': cliente.ciudad or 'N/A',
                '{{PROVINCIA}}': cliente.provincia or 'N/A',
                '{{DIRECCION_EXACTA}}': cliente.direccion_exacta or cliente.direccion or 'N/A',
                '{{DIRECCION}}': cliente.direccion_exacta or cliente.direccion or 'N/A',
                
                # Valores monetarios en números
                '{{VALOR_NUMEROS}}': f"{valor_sin_iva:.2f}",
                '{{IVA_NUMEROS}}': f"{valor_iva:.2f}",
                '{{TOTAL_NUMEROS}}': f"{valor_total:.2f}",
                
                # Valores monetarios en letras
                '{{VALOR_LETRAS}}': numero_a_letras(valor_sin_iva),
                '{{IVA_LETRAS}}': numero_a_letras(valor_iva),
                '{{TOTAL_LETRAS}}': numero_a_letras(valor_total),
                
                # Fechas
                '{{FECHA_INICIO}}': cuña.fecha_inicio.strftime('%d de %B del %Y') if cuña.fecha_inicio else 'N/A',
                '{{FECHA_FIN}}': cuña.fecha_fin.strftime('%d de %B del %Y') if cuña.fecha_fin else 'N/A',
                '{{FECHA_ACTUAL}}': timezone.now().strftime('%d de %B del %Y'),
                
                # Información del contrato
                '{{DURACION_DIAS}}': str(duracion_dias),
                '{{DURACION_MESES}}': str(duracion_meses),
                '{{SPOTS_DIA}}': str(cuña.repeticiones_dia),
                '{{DURACION_SPOT}}': str(cuña.duracion_planeada),
                '{{NUMERO_CONTRATO}}': self.numero_contrato,
                
                # Información adicional
                '{{CODIGO_CUÑA}}': cuña.codigo,
                '{{TITULO_CUÑA}}': cuña.titulo,
            }
            
            # Guardar datos de generación
            self.datos_generacion = datos_reemplazo
            
            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                for key, value in datos_reemplazo.items():
                    if key in paragraph.text:
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in datos_reemplazo.items():
                                if key in paragraph.text:
                                    for run in paragraph.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, value)
            
            # Guardar el documento modificado
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Guardar en el modelo
            from django.core.files.base import ContentFile
            nombre_archivo = f"contrato_{self.numero_contrato}.docx"
            self.archivo_contrato.save(
                nombre_archivo,
                ContentFile(buffer.read()),
                save=False
            )
            
            self.estado = 'generado'
            self.save()
            
            return True
            
        except Exception as e:
            print(f"Error generando contrato: {e}")
            return False
    
    def marcar_como_enviado(self):
        """Marca el contrato como enviado"""
        self.estado = 'enviado'
        self.fecha_envio = timezone.now()
        self.save()
    
    def marcar_como_firmado(self):
        """Marca el contrato como firmado"""
        self.estado = 'firmado'
        self.fecha_firma = timezone.now()
        self.save()
    
    def activar_contrato(self):
        """Activa el contrato"""
        if self.estado == 'firmado':
            self.estado = 'activo'
            self.save()
    
    def get_absolute_url(self):
        return reverse('content:contrato_detail', kwargs={'pk': self.pk})
    
    @property
    def puede_regenerar(self):
        """Verifica si el contrato puede regenerarse"""
        return self.estado in ['borrador', 'generado']
    
    @property
    def esta_activo(self):
        """Verifica si el contrato está activo"""
        return self.estado == 'activo' and self.cuña.esta_activa


# ==================== SEÑALES ====================

from django.db.models.signals import post_save, pre_save
from django.dispatch import receiver

@receiver(pre_save, sender=CuñaPublicitaria)
def cuña_pre_save(sender, instance, **kwargs):
    """Captura el estado anterior antes de guardar"""
    if instance.pk:
        try:
            instance._estado_anterior = CuñaPublicitaria.objects.get(pk=instance.pk)
        except CuñaPublicitaria.DoesNotExist:
            instance._estado_anterior = None
    else:
        instance._estado_anterior = None


@receiver(post_save, sender=CuñaPublicitaria)
def cuña_post_save(sender, instance, created, **kwargs):
    """Crea entrada en el historial después de guardar"""
    if created:
        HistorialCuña.objects.create(
            cuña=instance,
            accion='creada',
            usuario=getattr(instance, 'created_by', None),
            descripcion=f'Cuña publicitaria "{instance.titulo}" creada',
            datos_nuevos={
                'titulo': instance.titulo,
                'estado': instance.estado,
                'precio_total': str(instance.precio_total),
                'fecha_inicio': instance.fecha_inicio.isoformat() if instance.fecha_inicio else None,
                'fecha_fin': instance.fecha_fin.isoformat() if instance.fecha_fin else None,
            }
        )
    else:
        estado_anterior = getattr(instance, '_estado_anterior', None)
        if estado_anterior:
            if estado_anterior.estado != instance.estado:
                if instance.estado == 'aprobada':
                    HistorialCuña.objects.create(
                        cuña=instance,
                        accion='aprobada',
                        usuario=instance.aprobada_por,
                        descripcion=f'Cuña aprobada por {instance.aprobada_por}',
                        datos_anteriores={'estado': estado_anterior.estado},
                        datos_nuevos={'estado': instance.estado, 'aprobada_por': instance.aprobada_por.username if instance.aprobada_por else None}
                    )
            
            if estado_anterior.archivo_audio != instance.archivo_audio:
                accion = 'audio_subido' if not estado_anterior.archivo_audio else 'audio_cambiado'
                HistorialCuña.objects.create(
                    cuña=instance,
                    accion=accion,
                    usuario=getattr(instance, '_user_modificador', None),
                    descripcion=f'Archivo de audio {"subido" if accion == "audio_subido" else "cambiado"}',
                    datos_anteriores={'archivo_audio': str(estado_anterior.archivo_audio) if estado_anterior.archivo_audio else None},
                    datos_nuevos={'archivo_audio': str(instance.archivo_audio) if instance.archivo_audio else None}
                )
            
            if (estado_anterior.titulo != instance.titulo or 
                estado_anterior.precio_total != instance.precio_total or
                estado_anterior.fecha_inicio != instance.fecha_inicio or
                estado_anterior.fecha_fin != instance.fecha_fin):
                
                HistorialCuña.objects.create(
                    cuña=instance,
                    accion='editada',
                    usuario=getattr(instance, '_user_modificador', None),
                    descripcion='Cuña publicitaria editada',
                    datos_anteriores={
                        'titulo': estado_anterior.titulo,
                        'precio_total': str(estado_anterior.precio_total),
                        'fecha_inicio': estado_anterior.fecha_inicio.isoformat() if estado_anterior.fecha_inicio else None,
                        'fecha_fin': estado_anterior.fecha_fin.isoformat() if estado_anterior.fecha_fin else None,
                    },
                    datos_nuevos={
                        'titulo': instance.titulo,
                        'precio_total': str(instance.precio_total),
                        'fecha_inicio': instance.fecha_inicio.isoformat() if instance.fecha_inicio else None,
                        'fecha_fin': instance.fecha_fin.isoformat() if instance.fecha_fin else None,
                    }
                )
